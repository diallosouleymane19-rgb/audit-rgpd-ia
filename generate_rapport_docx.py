"""
generate_rapport_docx.py
Generateur de rapport micro-audit RGPD + AI Act au format Word (.docx)
Utilise python-docx — compatible Streamlit Cloud

Usage dans Streamlit :
    from generate_rapport_docx import generer_rapport
    buf = generer_rapport(client_info, score, nb_questions, niveau, ia_result)
    st.download_button("Telecharger le rapport .docx", buf, file_name="Rapport-Audit-....docx")
"""

from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Couleurs palette SMD ─────────────────────────────────────
NAVY   = RGBColor(0x1C, 0x29, 0x51)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
TEAL   = RGBColor(0x00, 0xB4, 0xA0)
RED    = RGBColor(0xEF, 0x44, 0x44)
AMBER  = RGBColor(0xF5, 0x9E, 0x0B)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
LGRAY  = RGBColor(0xE2, 0xE8, 0xF0)
DGRAY  = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

FONT_MAIN = "Arial"


# ─── Helpers ─────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Definit la couleur de fond d'une cellule (hex sans #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT_MAIN
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _heading(doc, text, level=1, color=NAVY):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    # Ligne de couleur sous le titre H1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "1C2951")
        bottom.set(qn("w:space"), "4")
        pBdr.append(bottom)
        pPr.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_MAIN
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = color
    return p


def _niveau_color(niveau: str):
    mapping = {
        "EXCELLENT":    ("16A34A", GREEN),
        "MOYEN":        ("F59E0B", AMBER),
        "INSUFFISANT":  ("F97316", RGBColor(0xF9, 0x73, 0x16)),
        "NON CONFORME": ("EF4444", RED),
    }
    return mapping.get(niveau, ("64748B", DGRAY))


def _prio_color(prio: str):
    return {
        "CRITIQUE":   ("EF4444", RED),
        "IMPORTANTE": ("F59E0B", AMBER),
        "INFO":       ("2563EB", BLUE),
    }.get(prio, ("64748B", DGRAY))


# ─── Generateur principal ─────────────────────────────────────

def generer_rapport(
    client_info: dict,   # {nom, contact, email, secteur, salaries}
    score: int,
    nb_questions: int,
    niveau: str,
    ia_result: dict,
) -> BytesIO:
    """
    Genere le rapport Word et retourne un BytesIO pret pour st.download_button.
    """
    doc = Document()

    # ── Marges ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Style par defaut ────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_MAIN
    style.font.size = Pt(10)

    nom_client = client_info.get("nom", "Client")
    date_str   = date.today().strftime("%d/%m/%Y")
    hex_niv, rgb_niv = _niveau_color(niveau)
    pct = round(score / nb_questions * 100) if nb_questions else 0

    # ══════════════════════════════════════════════════════════
    # PAGE DE GARDE
    # ══════════════════════════════════════════════════════════

    # Bandeau titre en couleur (simulé avec tableau 1 cellule)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, "1C2951")
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("RAPPORT MICRO-AUDIT")
    run.bold = True
    run.font.name = FONT_MAIN
    run.font.size = Pt(22)
    run.font.color.rgb = WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(18)
    r2 = p2.add_run("RGPD + EU AI Act")
    r2.bold = True
    r2.font.name = FONT_MAIN
    r2.font.size = Pt(14)
    r2.font.color.rgb = TEAL

    doc.add_paragraph()

    # Infos client
    _para(doc, nom_client, bold=True, size=20, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
    _para(doc, "Audit realise par SMD Global Consulting LLC", size=10, color=DGRAY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _para(doc, date_str, size=10, color=DGRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # Score visuel
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        ("Score", str(score) + " / " + str(nb_questions), "1C2951"),
        ("Niveau", niveau, hex_niv),
        ("Conformite", str(pct) + "%", hex_niv),
    ]
    for i, (lbl, val, hex_c) in enumerate(labels):
        c = tbl2.rows[0].cells[i]
        _set_cell_bg(c, hex_c)
        _set_cell_borders(c, hex_c)
        c.width = Cm(4.5)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(val)
        run.bold = True
        run.font.name = FONT_MAIN
        run.font.size = Pt(16)
        run.font.color.rgb = WHITE
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(lbl)
        r2.font.name = FONT_MAIN
        r2.font.size = Pt(9)
        r2.font.color.rgb = WHITE

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 1 — SYNTHESE EXECUTIVE
    # ══════════════════════════════════════════════════════════
    _heading(doc, "1. Synthese executive", level=1)

    synthese = ia_result.get("synthese_executive", "")
    if synthese:
        _para(doc, synthese, size=10, space_after=8)

    # Encadre risque amende
    risque = ia_result.get("risque_amende", "")
    if risque:
        tbl3 = doc.add_table(rows=1, cols=1)
        tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell3 = tbl3.rows[0].cells[0]
        _set_cell_bg(cell3, "FEF3C7")  # Amber clair
        _set_cell_borders(cell3, "F59E0B", "6")
        p_r = cell3.paragraphs[0]
        p_r.paragraph_format.space_before = Pt(6)
        r_label = p_r.add_run("Risque amende estime  ")
        r_label.bold = True
        r_label.font.name = FONT_MAIN
        r_label.font.size = Pt(10)
        r_label.font.color.rgb = AMBER
        p_r2 = cell3.add_paragraph()
        p_r2.paragraph_format.space_after = Pt(6)
        r_body = p_r2.add_run(risque)
        r_body.font.name = FONT_MAIN
        r_body.font.size = Pt(10)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 2 — RESULTATS PAR BLOC
    # ══════════════════════════════════════════════════════════
    _heading(doc, "2. Resultats par bloc", level=1)

    blocs = ia_result.get("analyse_blocs", {})
    if blocs:
        tbl4 = doc.add_table(rows=1, cols=4)
        tbl4.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl4.style = "Table Grid"

        # En-tete
        hdrs = ["Bloc", "Score", "Niveau", "Analyse"]
        widths = [Cm(1.5), Cm(2.2), Cm(2.8), Cm(9.5)]
        for i, (h, w) in enumerate(zip(hdrs, widths)):
            c = tbl4.rows[0].cells[i]
            _set_cell_bg(c, "1C2951")
            c.width = w
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(h)
            run.bold = True
            run.font.name = FONT_MAIN
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE

        for lettre, data in blocs.items():
            sb  = data.get("score", 0)
            mb  = data.get("max", 1)
            tb  = data.get("titre", "Bloc " + lettre)
            an  = data.get("analyse", "")
            pct_b = sb / mb if mb else 0
            if pct_b == 1.0:
                niv_b, hex_b = "Conforme",    "16A34A"
            elif pct_b >= 0.5:
                niv_b, hex_b = "Partiel",     "F59E0B"
            else:
                niv_b, hex_b = "Non conforme","EF4444"

            row = tbl4.add_row()
            vals = [lettre, str(sb) + "/" + str(mb), niv_b, an]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            bolds  = [True, True, True, False]
            for i, (v, al, bd) in enumerate(zip(vals, aligns, bolds)):
                c = row.cells[i]
                c.width = widths[i]
                if i == 2:
                    _set_cell_bg(c, hex_b)
                _set_cell_borders(c, "E2E8F0", "4")
                p = c.paragraphs[0]
                p.alignment = al
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                run = p.add_run(v)
                run.bold = bd
                run.font.name = FONT_MAIN
                run.font.size = Pt(9)
                if i == 2:
                    run.font.color.rgb = WHITE

    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 3 — RECOMMANDATIONS
    # ══════════════════════════════════════════════════════════
    _heading(doc, "3. Recommandations prioritaires", level=1)

    recos = ia_result.get("recommandations", [])
    for i, reco in enumerate(recos, 1):
        prio    = reco.get("priorite", "INFO")
        code    = reco.get("code_question", "")
        action  = reco.get("action", "")
        pourquoi = reco.get("pourquoi", "")
        delai   = reco.get("delai", "")
        article = reco.get("article", "")

        hex_p, rgb_p = _prio_color(prio)

        # Ligne titre reco
        tbl_r = doc.add_table(rows=1, cols=2)
        tbl_r.alignment = WD_TABLE_ALIGNMENT.LEFT
        c0 = tbl_r.rows[0].cells[0]
        c1 = tbl_r.rows[0].cells[1]
        _set_cell_bg(c0, hex_p)
        c0.width = Cm(3)
        c1.width = Cm(13)
        _set_cell_borders(c0, hex_p)
        _set_cell_borders(c1, "E2E8F0")
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after  = Pt(4)
        r0 = p0.add_run(str(i) + ". [" + prio + "]  " + code)
        r0.bold = True
        r0.font.name = FONT_MAIN
        r0.font.size = Pt(8)
        r0.font.color.rgb = WHITE

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after  = Pt(2)
        r1 = p1.add_run(action)
        r1.bold = True
        r1.font.name = FONT_MAIN
        r1.font.size = Pt(9)

        p2 = c1.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2a = p2.add_run("Pourquoi : ")
        r2a.bold = True
        r2a.font.size = Pt(8)
        r2a.font.name = FONT_MAIN
        r2a.font.color.rgb = DGRAY
        r2b = p2.add_run(pourquoi)
        r2b.font.size = Pt(8)
        r2b.font.name = FONT_MAIN

        p3 = c1.add_paragraph()
        p3.paragraph_format.space_after = Pt(6)
        r3a = p3.add_run("Delai : " + delai + "   |   Ref. : " + article)
        r3a.font.size = Pt(8)
        r3a.font.name = FONT_MAIN
        r3a.font.color.rgb = DGRAY

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 4 — PLAN D'ACTION
    # ══════════════════════════════════════════════════════════
    _heading(doc, "4. Plan d'action", level=1)

    plan = ia_result.get("plan_action", [])
    if plan:
        tbl5 = doc.add_table(rows=1, cols=6)
        tbl5.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl5.style = "Table Grid"
        plan_hdrs  = ["#", "Action", "Responsable", "Delai", "Ressource", "KPI"]
        plan_width = [Cm(0.8), Cm(5.5), Cm(2.5), Cm(2.0), Cm(2.5), Cm(2.7)]
        for i, (h, w) in enumerate(zip(plan_hdrs, plan_width)):
            c = tbl5.rows[0].cells[i]
            _set_cell_bg(c, "00B4A0")
            c.width = w
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(h)
            run.bold = True
            run.font.name = FONT_MAIN
            run.font.size = Pt(8)
            run.font.color.rgb = WHITE

        for item in plan:
            row = tbl5.add_row()
            vals = [
                str(item.get("ordre", "")),
                item.get("action", ""),
                item.get("responsable", ""),
                item.get("delai", ""),
                item.get("ressource", ""),
                item.get("kpi", ""),
            ]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER] + [WD_ALIGN_PARAGRAPH.LEFT] * 5
            for i, (v, al) in enumerate(zip(vals, aligns)):
                c = row.cells[i]
                c.width = plan_width[i]
                _set_cell_borders(c, "E2E8F0", "4")
                p = c.paragraphs[0]
                p.alignment = al
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                run = p.add_run(v)
                run.font.name = FONT_MAIN
                run.font.size = Pt(8)
                if i == 0:
                    run.bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # SECTION 5 — PROPOSITION COMMERCIALE
    # ══════════════════════════════════════════════════════════
    _heading(doc, "5. Proposition commerciale", level=1)

    opp = ia_result.get("opportunite_commerciale", "")
    if opp:
        tbl6 = doc.add_table(rows=1, cols=1)
        tbl6.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell6 = tbl6.rows[0].cells[0]
        _set_cell_bg(cell6, "EFF6FF")
        _set_cell_borders(cell6, "2563EB", "6")
        p6 = cell6.paragraphs[0]
        p6.paragraph_format.space_before = Pt(8)
        p6.paragraph_format.space_after  = Pt(8)
        r6 = p6.add_run(opp)
        r6.font.name = FONT_MAIN
        r6.font.size = Pt(10)

    doc.add_paragraph()

    # ── Grille tarifaire selon score ────────────────────────
    _heading(doc, "Grille tarifaire SMD", level=2)
    offres = [
        ("Score >= 90%",  "Veille reglementaire",              "150 EUR/mois",  "16A34A"),
        ("Score 65-90%",  "Mission conformite complete",        "1 500 EUR",     "F59E0B"),
        ("Score 40-65%",  "Mission urgente 6 mois",            "2 500 EUR",     "F97316"),
        ("Score < 40%",   "Mission complete + DPO externalise","4 000 EUR",     "EF4444"),
    ]
    tbl7 = doc.add_table(rows=len(offres), cols=3)
    tbl7.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, (seuil, mission, prix, hex_c) in enumerate(offres):
        cells = tbl7.rows[j].cells
        # Surligner la ligne correspondant au score actuel
        est_actif = (
            (pct >= 90  and j == 0) or
            (65 <= pct < 90 and j == 1) or
            (40 <= pct < 65 and j == 2) or
            (pct < 40  and j == 3)
        )
        bg = hex_c if est_actif else "F8FAFC"
        txt_color = WHITE if est_actif else None
        for k, (v, w) in enumerate(zip([seuil, mission, prix], [Cm(3.5), Cm(8), Cm(3)])):
            c = cells[k]
            _set_cell_bg(c, bg)
            _set_cell_borders(c, "E2E8F0", "4")
            c.width = w
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if k in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(v)
            run.font.name = FONT_MAIN
            run.font.size = Pt(9)
            run.bold = est_actif or k == 2
            if txt_color:
                run.font.color.rgb = WHITE

    doc.add_paragraph()

    # ── Pied de page ─────────────────────────────────────────
    _para(doc, "SMD Global Consulting LLC — diallosouleymane19@gmail.com",
          size=8, color=DGRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=2)
    _para(doc, "Document confidentiel — Prepare pour " + nom_client + " — " + date_str,
          size=8, color=DGRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    # ── Serialisation ────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
